"""DOCX renderer for the ABACUS Contract Governance Workbench."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .io import content_hash
from .schema import GovernanceSSOT

GENERATOR = "CODEX Contract Governance Generator"


def build_docx(payload: dict[str, object], ssot: GovernanceSSOT, path: Path) -> None:
    """Render a governance payload as a tiered DOCX snapshot."""

    document = Document()

    title = document.add_heading(
        f"{payload['package_id']} {str(payload['tier']).upper()} Governance Snapshot", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(f"Content hash (sha256): {content_hash(payload)}")

    for sheet_payload in payload["sheets"]:  # type: ignore[index]
        document.add_heading(sheet_payload["name"], level=2)
        columns = sheet_payload["columns"]
        rows = sheet_payload["rows"]
        table = document.add_table(rows=1 + len(rows), cols=len(columns))
        table.style = "Light Grid Accent 1"
        header_cells = table.rows[0].cells
        for index, column in enumerate(columns):
            header_cells[index].text = column
            for run in header_cells[index].paragraphs[0].runs:
                run.bold = True
        for row_index, row in enumerate(rows, start=1):
            cells = table.rows[row_index].cells
            for col_index, column in enumerate(columns):
                cells[col_index].text = str(row.get(column, ""))

    fixed = ssot.build.fixed_docprops_timestamp.replace(tzinfo=None)
    document.core_properties.created = fixed
    document.core_properties.modified = fixed
    document.core_properties.author = GENERATOR
    document.core_properties.last_modified_by = GENERATOR

    document.save(path)
