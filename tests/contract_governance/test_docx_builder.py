from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

from codex.contract_governance.builder import build_artifacts
from codex.contract_governance.io import load_ssot
from codex.contract_governance.validator import validate_generated

SSOT = Path("contract_governance/ssot/abacus_contract_governance.yaml")
FIXED_DOC_PROPS = datetime(2026, 1, 1, 0, 0, 0)
GENERATOR = "CODEX Contract Governance Generator"
FORBIDDEN_BIDDER_STRINGS = {
    "Extraction Audit",
    "Evaluation Notes",
    "REQ-W001-001",
    "Generated tier stripping",
    "Internal evaluation placeholder",
}


def _docx_full_text(path: Path) -> str:
    document = Document(path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join([paragraph_text, table_text])


def test_build_artifacts_generates_a_docx_snapshot(tmp_path: Path) -> None:
    ssot = load_ssot(SSOT)
    result = build_artifacts(ssot, tmp_path, "internal")

    assert "docx" in result
    docx_path = Path(result["docx"])
    assert docx_path.exists()

    document = Document(docx_path)
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style is not None and paragraph.style.name.startswith("Heading")
    ]
    assert any("Requirements" in heading for heading in headings)
    assert any("Extraction Audit" in heading for heading in headings)

    table_texts = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    assert "REQ-W001-001" in table_texts


def test_docx_pinned_core_properties_are_deterministic(tmp_path: Path) -> None:
    ssot = load_ssot(SSOT)
    first = build_artifacts(ssot, tmp_path / "first", "internal")
    second = build_artifacts(ssot, tmp_path / "second", "internal")

    for docx_path in (first["docx"], second["docx"]):
        document = Document(docx_path)
        # python-docx round-trips OPC core-properties dates as UTC-aware; openpyxl's
        # xlsx equivalent round-trips naive. Compare on the naive wall-clock value.
        assert document.core_properties.created.replace(tzinfo=None) == FIXED_DOC_PROPS
        assert document.core_properties.modified.replace(tzinfo=None) == FIXED_DOC_PROPS
        assert document.core_properties.author == GENERATOR
        assert document.core_properties.last_modified_by == GENERATOR


def test_bidder_docx_strips_internal_only_content(tmp_path: Path) -> None:
    ssot = load_ssot(SSOT)
    result = build_artifacts(ssot, tmp_path, "bidder")
    validate_generated(ssot, tmp_path, "bidder")

    docx_text = _docx_full_text(Path(result["docx"]))
    for forbidden in FORBIDDEN_BIDDER_STRINGS:
        assert forbidden not in docx_text
